import os
import sys
import docx2txt 
import re
import random
import docx
from docx import Document
from docx.shared import Pt

'''this library is necessary, you must pip install docx2txt #
in your command line so that we can read word documents
'''
folder_addr="/Users/kritikaravichander/qbpackettool/packets" #put the address of your folder with packets here
files=[]
sci=[]
hist=[]
lit=[]
fa=[]
rmpss=[]
geo=[]
ce=[]
misc=[]
trash=[]

sci_bonus=[]
hist_bonus=[]
lit_bonus=[]
fa_bonus=[]
rmpss_bonus=[]
geo_bonus=[]
ce_bonus=[]
misc_bonus=[]
trash_bonus=[]

cat_list=["SCI", "HIST", "LIT", "FA", "RMPSS", "GEO", "CE", "MISC.", "TRASH"]

#scop distro is
#5/5 Science (including 1/1 Noncomputational Math)
#4/4 Literature
#4/4 History
#2/2 Religion and Mythology
#2/2 Fine Arts
#1/1 Geography
#1/1 Philosophy, and Social Science
#1/1 trash (sports, pop cult)

dist={ #define the number of tossups in each category of the packet here
    "SCI":5,
    "HIST":4,
    "LIT":4,
    "FA":1,
    "RMPSS":1,
    "GEO":1,
    "CE":2,
    "MISC.":1,
    "TRASH":1
}

bonus_dist={ #define the number of bonuses in each category of the packet here
    "SCI":4,
    "HIST":5,
    "LIT":4,
    "FA":1,
    "RMPSS":1,
    "GEO":2,
    "CE":1,
    "MISC.":1,
    "TRASH":1
}

def getTossupList(category): #returns associated question list given tag
    if(category=="SCI"):
        return sci
    elif(category=="HIST"):
        return hist
    elif(category=="LIT"):
        return lit
    elif(category=="FA"):
        return fa
    elif(category=="RMPSS"):    
        return rmpss
    elif(category=="GEO"):
        return geo
    elif(category=="CE"):
        return ce
    elif(category=="MISC."):
        return misc
    elif(category=="TRASH"):
        return trash
    else:
        return []
def getBonusList(category): #returns associated bonus list given tag
    if(category=="SCI"):
        return sci_bonus
    elif(category=="HIST"):
        return hist_bonus
    elif(category=="LIT"):
        return lit_bonus
    elif(category=="FA"):
        return fa_bonus
    elif(category=="RMPSS"):
        return rmpss_bonus
    elif(category=="GEO"):
        return geo_bonus
    elif(category=="CE"):
        return ce_bonus
    elif(category=="MISC."):
        return misc_bonus
    elif(category=="TRASH"):
        return trash_bonus
    else:
        return []

def getCategory(question): #returns category name given the question test
    category=re.findall(r"<[^<\,]*\,", question)
    if(len(category)==0):
        return ""
    return (category[0][1:-1])

#tossup_dict={"SCI":sci, "HIST":hist, "LIT":lit, "FA":fa, "RMPSS":rmpss, "GEO":geo, "CE":ce, "MISC.":misc, "TRASH":trash}
#bonus_dict={"SCI":sci_bonus, "HIST":hist_bonus, "LIT":lit_bonus, "FA":fa_bonus, "RMPSS":rmpss_bonus, "GEO":geo_bonus, "CE":ce_bonus, "MISC.":misc_bonus, "TRASH":trash_bonus}

for file in os.listdir(folder_addr):
    if file.endswith('.docx'):
        files.append(file)
for i in range(len(files)): 
    file_addr=folder_addr+"/"+files[i]
    #print(len(files))
    text=docx2txt.process(file_addr)
    print(file_addr)
    texts=text.split("Bonuses")
    #print(texts)
    #text="1. dasdsklkdlsamds <HIST, WORLD> Bonus: Guyana <GEO, SA> 2. Kritika <SCI, CHEM> Bonus: Simping <RMPSS, PHIL>"
    all_tossups=re.findall(r"\d\.[^>]*>", texts[0])
    for x in range(len(all_tossups)):
        all_tossups[x]=all_tossups[x][3:]
    all_bonuses=re.findall(r"\d\.[^>]*>", texts[1])
    #print("Tossups: ", all_tossups[0] , "\n")
    #print("Bonuses: ", all_bonuses[0])
    for j in range(len(all_tossups)):
       # print(all_tossups[j])
        category=[]
        category=re.findall(r"<[^<\,]*\,", all_tossups[j])
        for x in range(len(category)):
            category[x]=category[x][1:-1]
            #print(category[x])
        (getTossupList(category[0])).append(all_tossups[j])
    for k in range(len(all_bonuses)):
        category=re.findall(r"<[^<\,]*\," , all_bonuses[k])
        for x in range(len(category)):
            category[x]=category[x][1:-1]
            #print(category[x]) 
        #print(category[0])
        #print("hello")  
        (getBonusList(category[0])).append(all_bonuses[k])
#print("Literature: ", lit)
#print("Geography: ", geo)


signal=1
packetnum=1
while signal!=0:
    packetList=[]
    packetBonuses=[]
    for cat in cat_list: #checks whether there are enough questions from each category to make a packet
        if(len(getTossupList(cat))<dist[cat]):
            signal=0
        if(len(getBonusList(cat))<bonus_dist[cat]):
            signal=0
    if(signal==1):
        for cat in cat_list:
            for x in range(dist[cat]):
                question_num=random.randint(0, len(getTossupList(cat))-1)
                packetList.append(((getTossupList(cat))[question_num]))
                getTossupList(cat).pop(question_num)
            for y in range(bonus_dist[cat]):
                question_num=random.randint(0, len(getBonusList(cat))-1)
                packetBonuses.append(((getBonusList(cat))[question_num]))
                getBonusList(cat).pop(question_num)
        val=0
        count=0
        while(val!=1):
            val=1
            random.shuffle(packetList)
            for x in range(1, len(packetList)): 
                if(getCategory(packetList[x])==getCategory(packetList[x-1])):
                    val=0
                    break
            count=count+1
            if(count==100):
                break
        val=0
        count=0
        while(val!=1):
            val=1
            random.shuffle(packetBonuses)
            for x in range(1, len(packetBonuses)): 
                if(getCategory(packetBonuses[x])==getCategory(packetBonuses[x-1])):
                    val=0
                    break
            count=count+1
            if(count==100):
                break
                
        #print("Tossups", packetList)
        #print("Bonuses", packetBonuses)
        doc=docx.Document()
        header="QQBC Packet "+str(packetnum)
        doc.add_heading(header, 0)
        style=doc.styles['Normal']
        font=style.font
        font.name='Arial'
        font.size=Pt(10)

        for i in range(min(len(packetList), len(packetBonuses))):
            tossup_text=""+str(i+1)
            tossup_para=doc.add_paragraph(tossup_text)
            parts=packetList[i].split("(*)")
            tossup_para.add_run(". ")
            tossup_para.add_run(parts[0]+"(*)").bold = True
            tossup_para.add_run(parts[1])
            bonus_text=packetBonuses[i]
            bonus_para=doc.add_paragraph(bonus_text)
            path='/Users/kritikaravichander/qbpackettool/generated_packets/'+'QQBC_Packet'+str(packetnum)+'.docx'
            doc.save(path)

    packetnum=packetnum+1
packetnum=packetnum-1
print(packetnum-1, " packets completed")


#extra packets
packetList=[]
packetBonuses=[]
for cat in cat_list:
    for t in getTossupList(cat):
        packetList.append(t)
    for b in getBonusList(cat):
        packetBonuses.append(b)
val=0
count=0
while(val!=1):
    val=1
    random.shuffle(packetList)
    for x in range(1, len(packetList)): 
        if(getCategory(packetList[x])==getCategory(packetList[x-1])):
            val=0
            break
    count=count+1
    if(count==100):
        break
val=0
count=0
while(val!=1):
    val=1
    random.shuffle(packetBonuses)
    for x in range(1, len(packetBonuses)): 
        if(getCategory(packetBonuses[x])==getCategory(packetBonuses[x-1])):
            val=0
            break
    count=count+1
    if(count==100):
        break
print(len(packetList))
print(len(packetBonuses))
while len(packetList)>=20 and len(packetBonuses)>=20:
    doc=docx.Document()
    header="QQBC Packet "+str(packetnum)
    doc.add_heading(header, 0)
    style=doc.styles['Normal']
    font=style.font
    font.name='Arial'
    font.size=Pt(10)

    for i in range(20):
        tossup_text=""+str(i+1)
        tossup_para=doc.add_paragraph(tossup_text)
        parts=packetList[i].split("(*)")
        tossup_para.add_run(". ")
        tossup_para.add_run(parts[0]+"(*)").bold = True
        tossup_para.add_run(parts[1])
        bonus_text=packetBonuses[i]
        bonus_para=doc.add_paragraph(bonus_text)
        path='/Users/kritikaravichander/qbpackettool/generated_packets/'+'QQBC_Packet'+str(packetnum)+'.docx'
        doc.save(path)
    for i in range(20):
        packetList.remove(packetList[19-i])
        packetBonuses.remove(packetBonuses[19-i])

    packetnum=packetnum+1


    
